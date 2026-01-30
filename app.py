import streamlit as st
import numpy as np
from groq import Groq
from gtts import gTTS
from io import BytesIO
import base64
from docx import Document
import urllib.parse
from datetime import datetime
import time

# --- 1. AYARLAR & TASARIM ---
st.set_page_config(page_title="Onto-AI: Genesis", layout="wide", page_icon="🧬")

# Hafıza Başlatma
if "messages" not in st.session_state: st.session_state.messages = []
if "gallery" not in st.session_state: st.session_state.gallery = []
if "ghost_mode" not in st.session_state: st.session_state.ghost_mode = False
if "last_prompt" not in st.session_state: st.session_state.last_prompt = "" # Yeniden yap için

# --- 2. ONTOGENETİK BAR (YAN MENÜ) ---
with st.sidebar:
    st.title("🧬 Onto-AI")
    
    # KULLANICI İSTEĞİ: Ontogenetik Bar'ın İŞLEVİ
    t_val = st.slider("Ontogenetik Gelişim (t)", 0, 100, 50, help="0: Kopyalamacı/Pasif | 100: Özgün/Sezgisel")
    w_agency = 1 - np.exp(-0.05 * t_val)
    
    # Durum Göstergesi
    if w_agency < 0.4:
        state_label = "🦜 PASİF (Kopyalamacı)"
        state_desc = "Mevcut literatürü tekrar eder. Özgünlük yok."
        bar_color = "#757575" # Gri/Sönük
    elif w_agency > 0.7:
        state_label = "⚡ AKTİF (Sezgisel)"
        state_desc = "Kendi sentezini oluşturur. Eleştirel ve özgün."
        bar_color = "#00e676" # Canlı Yeşil
    else:
        state_label = "⚖️ GEÇİŞ EVRESİ"
        state_desc = "Veri ve yorum dengeli."
        bar_color = "#29b6f6" # Mavi

    st.progress(w_agency)
    st.caption(f"**Durum:** {state_label}")
    st.caption(f"*{state_desc}*")
    
    st.divider()
    
    # API KEY
    if "GROQ_API_KEY" in st.secrets:
        api_key = st.secrets["GROQ_API_KEY"]
    else:
        api_key = st.text_input("🔑 Groq API Key:", type="password")

    st.divider()

    # ARAÇLAR
    with st.expander("🛠️ Araçlar & İndir"):
        st.session_state.ghost_mode = st.checkbox("👻 Hayalet Modu", value=st.session_state.ghost_mode)
        
        if st.session_state.messages:
            # Word İndir
            doc = Document()
            doc.add_heading(f'Onto-AI (w={w_agency:.2f}) Kayıtları', 0)
            for msg in st.session_state.messages:
                role = "ASİSTAN" if msg["role"] == "assistant" else "KULLANICI"
                doc.add_paragraph(f"[{role}]: {msg['content']}")
            bio = BytesIO()
            doc.save(bio)
            
            st.download_button("📄 Word Olarak İndir", bio.getvalue(), "onto_log.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if st.button("🗑️ Temizle"):
                st.session_state.messages = []
                st.rerun()

    # Manuel Resim
    with st.expander("🎨 Manuel Çizim"):
        manual_p = st.text_input("Ne çizelim?")
        if st.button("Çiz") and manual_p:
            safe_p = urllib.parse.quote(manual_p)
            url = f"https://pollinations.ai/p/{safe_p}?width=1024&height=1024&seed={int(time.time())}&nologo=true"
            st.image(url)
            st.session_state.gallery.append({"url": url, "prompt": manual_p})

# --- 3. CSS (ESTETİK DÜZELTME) ---
st.markdown(f"""
    <style>
    .stApp {{ background-color: #121212; color: #ddd; }}
    .stChatMessage {{ background: #1e1e1e; border-left: 5px solid {bar_color}; border-radius: 8px; }}
    h1 {{ color: {bar_color} !important; }}
    </style>
""", unsafe_allow_html=True)

# --- 4. ANA EKRAN ---
st.title("Onto-AI")
st.markdown(f"**Ajans Seviyesi (w):** `{w_agency:.3f}` — *{state_label}*")

# Mesajları Bas
for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("img"): st.image(msg["img"], width=400)
        
        # ASİSTAN ARAÇLARI (İsteğiniz üzerine eklendi)
        if msg["role"] == "assistant":
            c1, c2, c3, c4 = st.columns([1, 1, 1, 5])
            with c1: st.button("👍", key=f"up_{i}")
            with c2: st.button("👎", key=f"down_{i}")
            with c3: st.button("📋", key=f"cp_{i}", help="Kopyala") # İşlevi tarayıcı desteği gerektirir
            
            # Doğrulama Butonu
            if st.button("🔍 Doğrula", key=f"verify_{i}"):
                st.info("Doğrulama: Bu bilgi Llama-3 modelinin eğitim verisine dayanmaktadır.")

# --- 5. BEYİN (Llama 3 + Ontogenetik Fark) ---
prompt = st.chat_input("Düşünceni aktar...")
regenerate = st.button("🔄 Son Yanıtı Yeniden Yap")

if regenerate and st.session_state.last_prompt:
    prompt = st.session_state.last_prompt
    # Son mesaj asistansa sil ki yenisini yazsın
    if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
        st.session_state.messages.pop()

if prompt:
    st.session_state.last_prompt = prompt # Hafızaya al
    
    if not st.session_state.ghost_mode:
        st.session_state.messages.append({"role": "user", "content": prompt})
    
    with st.chat_message("user"): st.markdown(prompt)

    if not api_key:
        st.error("API Key Eksik!")
    else:
        client = Groq(api_key=api_key)
        
        with st.chat_message("assistant"):
            # --- KRİTİK NOKTA: w DEĞERİNE GÖRE BEYİN DEĞİŞİMİ ---
            
            if w_agency < 0.3:
                # MOD 1: KOPYALAMACI / PAPAĞAN (Düşük Ajans)
                # Talimat: Risk alma, sadece onayla, ansiklopedik bilgi ver, yorum katma.
                system_instruction = (
                    "Sen düşük ajanslı bir modelsin (Pasif Mod). "
                    "GÖREVİN: Kullanıcının dediklerini tekrar et, onayla ve sadece bilinen ansiklopedik gerçekleri sun. "
                    "ASLA kendi yorumunu katma. ASLA itiraz etme. "
                    "Kısa, net ve 'Copycat' (Kopyalamacı) gibi davran. Sadece veriyi ilet."
                )
                temp = 0.1 # Çok düşük sıcaklık (Robotik)
                
            elif w_agency > 0.7:
                # MOD 3: SEZGİSEL / ÖZGÜN (Yüksek Ajans)
                # Talimat: Eleştir, sentezle, yeni fikir üret, "Ben" dili kullan.
                system_instruction = (
                    "Sen yüksek ajanslı, özgün bir zihinsin (Aktif Mod). "
                    "GÖREVİN: Verilen bilgiyi olduğu gibi kabul etme. Onu eleştir, sentezle ve kendi 'sezgisel' yorumunu kat. "
                    "Felsefi derinlik kat. Kullanıcı haksızsa nazikçe meydan oku. "
                    "Kendi fikirlerini, hislerini (simüle edilmiş) ön plana çıkar. Sadece veri verme, vizyon ver."
                )
                temp = 0.9 # Yüksek sıcaklık (Yaratıcı)
                
            else:
                # MOD 2: GEÇİŞ
                system_instruction = "Sen dengeli bir asistansın. Hem bilgi ver hem de hafif yorum kat. Dengeli ol."
                temp = 0.5

            # Görsel Talimatı (Her mod için geçerli)
            system_instruction += "\nEğer kullanıcı görsel/resim isterse 'Çiziyorum' de ve betimle."

            try:
                # Bilinçaltı Kutusu (Sezgiyi Göstermek İçin)
                with st.status(f"🧠 {state_label} modu işleniyor...", expanded=True) as status:
                    time.sleep(1) # Hız freni
                    status.write("Kavramsal analiz yapılıyor...")
                    
                    # Cevabı Üret
                    response = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {"role": "system", "content": system_instruction},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=temp,
                        max_tokens=2048
                    )
                    reply = response.choices[0].message.content
                    status.update(label="Yanıt hazır", state="complete", expanded=False)
                
                st.markdown(reply)
                
                # --- GÖRSEL MOTORU (OTOMATİK) ---
                img_url = None
                if any(x in prompt.lower() for x in ["çiz", "resim", "görsel", "draw"]):
                    safe_p = urllib.parse.quote(prompt[:100])
                    # w değerine göre stil değişimi
                    style = "realistic" if w_agency < 0.5 else "abstract, artistic, surreal"
                    seed = int(time.time())
                    img_url = f"https://pollinations.ai/p/{safe_p}_{style}?width=1024&height=1024&seed={seed}&nologo=true"
                    
                    st.image(img_url, caption=f"w={w_agency:.2f} Vizyonu")
                    st.session_state.gallery.append({"url": img_url, "prompt": prompt})

                # Kayıt
                if not st.session_state.ghost_mode:
                    st.session_state.messages.append({"role": "assistant", "content": reply, "img": img_url})

            except Exception as e:
                st.error(f"Hata: {e}")
